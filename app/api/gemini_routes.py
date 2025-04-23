from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query, Path as PathParam, Form, BackgroundTasks, Request, Header
from fastapi.responses import HTMLResponse
from sqlalchemy.orm import Session
from typing import List, Dict, Any, Optional, Union
import tempfile
import os
import PyPDF2
import textwrap
from io import BytesIO
import asyncio
from sqlalchemy import text
from datetime import datetime
import json
from sqlalchemy import inspect
from sqlalchemy.exc import SQLAlchemyError
import traceback
import subprocess
import sys
from pathlib import Path
import re
import uuid

from app.db.database import get_db, engine
from app.models.api_models import (
    CompletionRequest, CompletionResponse, 
    EmbeddingRequest, EmbeddingResponse,
    DocumentCreate, DocumentResponse,
    QueryRequest, QueryResponse
)
from app.services.gemini_service import GeminiService
from app.services.vector_service import VectorService
from app.services.db_service import DatabaseService

# Create a unified router, no longer need separate authenticated and non-authenticated routes
router = APIRouter(prefix="/api/v1")
health_router = APIRouter(prefix="/api/v1")

# Create service instances
gemini_service = GeminiService()

# 创建DatabaseService的临时实例，用于初始化VectorService
def get_vector_service(db: Session = Depends(get_db)):
    db_service = DatabaseService(db)
    return VectorService(db_service, gemini_service)

# 使用全局变量保存VectorService实例，在路由中使用依赖注入
vector_service = None

# Health check endpoint - no authentication required
@health_router.get("/health", 
                  summary="API Health Check", 
                  description="Check if the API service is running properly, no authentication required",
                  response_description="Returns the health status of the API service and the current timestamp")
async def health():
    """
    Health check endpoint to verify if the service is running properly
    
    No authentication required, can be used by monitoring systems to check API availability
    
    Returns:
        dict: Dictionary containing status and timestamp
    """
    return {"status": "ok", "timestamp": datetime.now().isoformat()}

# Database status endpoint - no authentication required
@health_router.get("/database-status", 
                  summary="Database Connection Status", 
                  description="Check database connection status, no authentication required",
                  response_description="Returns database connection status and current timestamp")
async def database_status():
    """
    Check database connection status
    
    No authentication required, can be used by monitoring systems to check API availability
    
    Returns:
        dict: Dictionary containing database connection status and timestamp
        
    Exceptions:
        HTTPException: If database connection fails
    """
    try:
        with engine.connect() as connection:
            result = connection.execute(text("SELECT 1"))
            if result:
                return {
                    "status": "connected",
                    "timestamp": datetime.now().isoformat()
                }
    except SQLAlchemyError as e:
        raise HTTPException(status_code=500, detail=f"Database connection failed: {str(e)}")

# All endpoints below do not require authentication

@router.post("/embedding", 
            response_model=EmbeddingResponse, 
            summary="Generate Text Embedding Vector", 
            description="Convert text to embedding vector representation for vector retrieval",
            response_description="Returns the generated vector representation")
async def create_embedding(
    request: EmbeddingRequest,
):
    """
    Generate embedding vectors for text
    
    Use Gemini embedding model to convert input text to vector representation for similarity search
    
    Parameters:
        request: Request object containing the text to be converted
        
    Returns:
        EmbeddingResponse: Response object containing the embedding vector
        
    Exceptions:
        HTTPException: If embedding generation fails
    """
    try:
        embedding = await gemini_service.generate_embedding(request.text)
        return {"embedding": embedding}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Embedding generation failed: {str(e)}")

@router.post("/completion", 
            response_model=CompletionResponse, 
            summary="Generate Text Completion", 
            description="Use Gemini model to generate text completion, optionally using relevant documents from the vector database as context",
            response_description="Returns the generated completion text")
async def create_completion(
    request: CompletionRequest,
    db: Session = Depends(get_db),
    vector_service: VectorService = Depends(get_vector_service),
):
    """
    Generate text completion
    
    Use Gemini model to generate text completion, optionally using relevant documents from the vector database as context
    
    Parameters:
        request: Request object containing prompt text and context configuration
        db: Database session
        vector_service: Vector service instance
        
    Returns:
        CompletionResponse: Response object containing the generated completion text
        
    Exceptions:
        HTTPException: If completion generation fails
    """
    try:
        context = None
        
        # If context is needed
        if request.use_context and request.context_query:
            # Get similar documents
            similar_docs = await vector_service.search_similar(
                db, 
                request.context_query, 
                request.max_context_docs
            )
            
            # Prepare context
            if similar_docs:
                context = await gemini_service.prepare_context(request.context_query, similar_docs)
        
        # Determine task complexity based on prompt length and content
        complexity = "normal"
        if request.model_complexity in ["simple", "normal", "complex"]:
            # Use user-specified complexity if provided
            complexity = request.model_complexity
        elif len(request.prompt) > 1000 or "详细" in request.prompt or "complex" in request.prompt.lower():
            complexity = "complex"
        elif len(request.prompt) < 100 and not any(term in request.prompt.lower() for term in ["explain", "analyze", "compare", "evaluate"]):
            complexity = "simple"
        
        # Generate completion with complexity hint
        completion = await gemini_service.generate_completion(
            prompt=request.prompt, 
            context=context,
            complexity=complexity,
            use_cache=not request.disable_cache
        )
        
        return {"completion": completion}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Completion generation failed: {str(e)}")

@router.post("/documents", 
            response_model=DocumentResponse, 
            summary="Add Document", 
            description="Add a single document to the vector database, generating its embedding vector for later retrieval",
            response_description="Returns information about the added document")
async def add_document(
    request: DocumentCreate,
    db: Session = Depends(get_db),
    vector_service: VectorService = Depends(get_vector_service),
):
    """
    Add document to the vector database
    
    Add a single document to the database, generating its embedding vector for later retrieval
    
    Parameters:
        request: Request object containing document content and metadata
        db: Database session
        vector_service: Vector service instance
        
    Returns:
        DocumentResponse: Response object containing information about the newly added document
        
    Exceptions:
        HTTPException: If adding the document fails
    """
    try:
        document = await vector_service.add_document(db, request.content, request.metadata)
        
        # 手动创建返回对象，确保符合DocumentResponse模型的要求
        # 从文档元数据中获取并解析元数据
        metadata = {}
        if document.doc_metadata:
            try:
                import json
                metadata_dict = json.loads(document.doc_metadata)
                # 排除内部字段如_embedding
                metadata = {k: v for k, v in metadata_dict.items() if not k.startswith('_')}
            except Exception as e:
                print(f"无法解析文档元数据: {e}")
        
        # 创建符合DocumentResponse的字典
        return {
            "id": document.id,
            "content": request.content,  # 使用请求中的内容，因为数据库中可能没有存储
            "metadata": metadata
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to add document: {str(e)}")

@router.post("/query", 
            response_model=QueryResponse, 
            summary="Query Similar Documents", 
            description="Search for documents similar to the query text in the vector database, and return detailed information",
            response_description="Returns a list of similar documents, context, and summary analysis")
async def query_documents(
    request: QueryRequest,
    source_filter: Optional[str] = Query(None, description="Filter results by document source", example="document1.pdf"),
    db: Session = Depends(get_db),
    vector_service: VectorService = Depends(get_vector_service),
):
    """
    Query similar documents in the vector database
    
    Search for documents similar to the query text in the vector database, and return detailed information, context, and summary analysis
    
    Parameters:
        request: Request object containing query text and configuration
        source_filter: Optional document source filter condition
        db: Database session
        vector_service: Vector service instance
        
    Returns:
        QueryResponse: Response object containing similar documents, context, and summary
        
    Exceptions:
        HTTPException: If the query fails
    """
    try:
        results = await vector_service.search_similar(db, request.query, request.limit, source_filter)
        return {"results": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to query similar documents: {str(e)}")

@router.post("/integration", 
            response_model=CompletionResponse, 
            summary="Integration Query", 
            description="Integrated API combining vector retrieval and text generation, can filter results by document source and return detailed debugging information",
            response_description="Returns answer generated based on related document content")
async def integration_query(
    request: CompletionRequest,
    source_filter: Optional[str] = Query(None, description="Filter results by document source", example="document1.pdf"),
    debug: bool = Query(False, description="Return detailed debugging information"),
    force_use_documents: bool = Query(False, description="Force use document content answer even if no high similarity match found"),
    db: Session = Depends(get_db),
    vector_service: VectorService = Depends(get_vector_service),
):
    """
    Integrated API: Vector query and result addition to prompt
    
    Integrated API combining vector retrieval and text generation, can filter results by document source and return detailed debugging information
    
    Parameters:
        request: Request object containing prompt text and context query
        source_filter: Filter results by document source, for example specific PDF file name
        debug: Whether to return detailed debugging information
        force_use_documents: Force use document content answer even if no high similarity match found
        db: Database session
        vector_service: Vector service instance
        
    Returns:
        CompletionResponse: Response object containing the generated answer, if debug enabled then includes debugging information
        
    Exceptions:
        HTTPException: If integration query fails
    """
    try:
        print(f"Received integration query request: {request.prompt}")
        print(f"Context query: {request.context_query or request.prompt}")
        
        # Check if it's a Chinese query
        is_chinese_query = any('\u4e00' <= c <= '\u9fff' for c in request.prompt)
        print(f"Is Chinese query: {is_chinese_query}")
        
        # 检测是否是表格相关查询
        is_table_query = any(term in request.prompt.lower() for term in [
            "table", "row", "column", "ranking", "largest", "smallest", "third", "3rd", "top", 
            "first", "second", "fourth", "fifth", "next", "previous", "highest", "lowest", 
            "percentage", "比例", "表格", "行", "列", "排名", "最大", "最小", "第三", "前",
            "country", "bond", "bonds", "holding", "investment", "government"
        ])
        
        # 检测更多数字排名相关术语
        has_ranking_terms = any(term in request.prompt.lower() for term in [
            "largest", "biggest", "highest", "top", "smallest", "lowest", "bottom",
            "first", "second", "third", "fourth", "fifth", "1st", "2nd", "3rd", "4th", "5th",
            "排名第", "最高的", "最低的", "第一", "第二", "第三", "第四", "第五"
        ])
        
        # 如果问题中包含排名相关术语，强制使用文档内容回答
        if has_ranking_terms:
            force_use_documents = True
            print(f"检测到排名相关查询，强制使用文档内容回答")
            
        print(f"Is table-related query: {is_table_query}")
        print(f"Force use documents: {force_use_documents}")
        
        # 增加表格相关的搜索 limit
        max_context_docs = request.max_context_docs
        if is_table_query:
            max_context_docs = max(max_context_docs, 20)  # 表格查询增加文档返回数量
            print(f"Increased max_context_docs to {max_context_docs} for table-related query")
        
        # First query related documents
        search_query = request.context_query or request.prompt
        
        # Expand search query to improve matching ability for Chinese content
        expanded_terms = []
        if "道教" in search_query:
            expanded_terms = ["道教", "老子", "道德经", "太上老君", "张道陵", "太极", "阴阳", "天师道", "五斗米道", "全真道"]  # Taoism, Laozi, Tao Te Ching, Supreme Old Lord, Zhang Daoling, Taiji, Yin-Yang, Celestial Masters, Five Pecks of Rice, Complete Perfection
        elif "佛教" in search_query:
            expanded_terms = ["佛教", "释迦牟尼", "佛陀", "菩萨", "禅宗", "佛经", "如来", "佛祖", "涅槃", "菩提"]  # Buddhism, Shakyamuni, Buddha, Bodhisattva, Zen, Buddhist Scriptures, Tathagata, Founder of Buddhism, Nirvana, Bodhi
        elif is_chinese_query:
            # Add some common Chinese philosophy and religious terms for other Chinese queries
            expanded_terms = ["中国", "哲学", "历史", "传统", "文化", "典籍", "经典"]  # China, philosophy, history, tradition, culture, ancient books, classics
        elif is_table_query:
            # 添加表格和财务相关的扩展术语
            if "bond" in search_query.lower() or "investment" in search_query.lower():
                expanded_terms = ["bond", "bonds", "holdings", "government", "sovereign", "treasury", "debt", 
                                  "portfolio", "investment", "securities", "fixed income", "yield", "maturity",
                                  "country", "countries", "allocation", "percentage", "asset", "table", "data"]
            else:
                expanded_terms = ["table", "data", "chart", "figure", "statistics", "number", "percentage", "ranking"]
        
        if expanded_terms:
            search_query = f"{search_query} {' '.join(expanded_terms)}"
            print(f"Expanded search query: {search_query}")
        
        # Increase return document count to improve probability of finding related content
        if is_chinese_query:
            max_context_docs = max(max_context_docs, 10)  # Chinese query at least returns 10 documents
        
        # 对表格相关查询，特别是排名类查询，增加相似度阈值，以确保得到最相关的文档
        similar_docs = await vector_service.search_similar(
            db, 
            search_query, 
            max_context_docs,
            source_filter
        )
        
        print(f"Found {len(similar_docs)} related documents")
        
        # 增加对文档相似度的更详细分析
        docs_with_high_similarity = [doc for doc in similar_docs if doc.get("similarity", 0) > 0.7]
        print(f"Found {len(docs_with_high_similarity)} documents with high similarity (>0.7)")
        
        # 针对表格查询，如果相似度不够高，强制使用文档内容
        if is_table_query and has_ranking_terms and not docs_with_high_similarity:
            force_use_documents = True
            print("针对表格排名查询，没有找到高相似度文档，强制使用文档内容")
        
        # For debugging purposes, print content snippets of the first few documents
        if similar_docs:
            print("First 3 document content snippets:")
            for i, doc in enumerate(similar_docs[:3]):
                content_preview = doc.get("content", "")[:100].replace("\n", " ")
                similarity = doc.get("similarity", 0)
                print(f"  [{i+1}] Similarity: {similarity:.4f} - {content_preview}...")
        
        # Prepare context
        context = None
        if similar_docs:
            context = await gemini_service.prepare_context(
                request.context_query or request.prompt, 
                similar_docs
            )
            print(f"Generated context, length: {len(context) if context else 0}")
        else:
            print("No related documents found")
        
        # Generate completion
        completion_prompt = request.prompt
        if context:
            # 添加表格处理特殊指令
            table_instruction = ""
            if is_table_query:
                table_instruction = """
IMPORTANT: The reference material may contain table data and structured information. When answering:
1. Pay special attention to any tables, data series, or numerical lists in the reference material
2. Maintain the structural relationships between entities in tables (rows, columns, rankings)
3. Verify numerical relationships carefully (largest, third-largest, percentages, etc.)
4. If you identify a table, first reconstruct it mentally to understand the data structure
5. Be precise when citing specific values, positions, or rankings from tabular data
"""
                
                # 对于涉及排名的问题，添加更强的指导
                if has_ranking_terms:
                    table_instruction += """
6. This question involves ranking or ordering. BE EXTREMELY CAREFUL to:
   - Identify the exact ranking criteria (largest by what measure?)
   - Count positions accurately when identifying rankings (1st, 2nd, 3rd, etc.)
   - Double-check your answer by reviewing the data in the reference material
   - Cite the specific table or data source that supports your answer
   - For third-largest, ensure you are identifying the item in the THIRD position, not any other position
"""
            
            if is_chinese_query:
                completion_prompt = f"""
You are a knowledgeable assistant tasked with answering questions based ONLY on the provided reference material.

IMPORTANT INSTRUCTIONS:
1. You are given English documents but the user is asking in Chinese.
2. You must answer in Chinese.
3. You must ONLY use information found in the reference material.
4. Do NOT use your general knowledge unless absolutely necessary to explain concepts in the documents.
5. If the reference material does not contain the answer, clearly state this in Chinese.
6. Translate relevant information from the English documents into Chinese to answer the query.
7. When referencing specific points from the documents, mention which document they came from.
{table_instruction}

Reference material:
{context}

User question (in Chinese): {request.prompt}

Provide a comprehensive answer in Chinese, based strictly on the information in the reference material.
"""
            else:
                completion_prompt = f"""
You are a knowledgeable assistant tasked with answering questions based ONLY on the provided reference material.

IMPORTANT INSTRUCTIONS:
1. You must ONLY use information found in the reference material.
2. Do NOT use your general knowledge unless absolutely necessary to explain concepts in the documents.
3. If the reference material does not contain the answer, clearly state this.
4. When referencing specific points from the documents, mention which document they came from.
{table_instruction}

Reference material:
{context}

User question: {request.prompt}

Provide a comprehensive answer based strictly on the information in the reference material.
"""
        else:
            # If no related documents found but user forces use of document content
            if force_use_documents:
                # Get existing documents in the system
                try:
                    # Query document count
                    doc_count = db.execute(text("SELECT COUNT(*) FROM documents")).scalar()
                    
                    # If there are documents, try to get some sample content
                    if doc_count > 0:
                        # Get recent document titles and sources
                        recent_docs = db.execute(
                            text("SELECT id, title, doc_metadata FROM documents ORDER BY id DESC LIMIT 5")
                        ).fetchall()
                        
                        doc_samples = []
                        for doc in recent_docs:
                            metadata = {}
                            if doc.doc_metadata:
                                try:
                                    if isinstance(doc.doc_metadata, str):
                                        metadata = json.loads(doc.doc_metadata)
                                    else:
                                        metadata = doc.doc_metadata
                                except:
                                    pass
                            
                            source = metadata.get("source", "Unknown source")
                            doc_samples.append(f"- ID: {doc.id}, Source: {source}, Title: {doc.title[:50]}...")
                        
                        doc_list = "\n".join(doc_samples)
                        
                        if is_chinese_query:
                            completion_prompt = f"""
非常抱歉，我在系统中搜索了所有文档（共 {doc_count} 个），但未能找到与您的问题"{request.prompt}"直接相关的内容。

以下是系统中的一些文档示例：
{doc_list}

由于您选择了强制使用文档内容模式，我必须基于系统中的文档回答问题，但系统中可能没有与此问题相关的信息。

请考虑以下几点：
1. 是否需要上传包含相关信息的文档
2. 是否需要调整查询措辞以更好地匹配现有文档
3. 是否需要扩展文档库以涵盖这个主题

您的问题是：{request.prompt}

我必须声明，系统中可能缺少相关文档。
"""
                        else:
                            completion_prompt = f"""
I am a knowledgeable assistant, especially skilled at answering questions about documents in your system.

I searched through all documents in the database ({doc_count} total) but couldn't find any directly relevant to your question about "{request.prompt}".

Here are some sample documents in the system:
{doc_list}

Since this query is using the force_use_documents mode, I must base my answer only on documents in the system. However, the system doesn't appear to contain information related to your specific question.

Please consider:
1. Uploading documents containing the relevant information about {request.prompt}
2. Adjusting your query wording to better match existing documents
3. Expanding the document library to cover this topic

I cannot provide a specific answer to your question as the necessary information doesn't appear to be in the document database.
"""
                except Exception as e:
                    print(f"Failed to get document statistics: {e}")
                    # Use default prompt
                    if is_chinese_query:
                        completion_prompt = "非常抱歉，我在文档库中找不到与您问题相关的信息。请上传包含相关数据的文档，以便我能提供准确的回答。"
                    else:
                        completion_prompt = "I'm sorry, I cannot find information related to your question in the document database. Please upload relevant documents to get an accurate answer."
            
            elif not force_use_documents:
                if is_chinese_query:
                    completion_prompt = f"""
非常抱歉，我在系统中没有找到与"{request.prompt}"相关的参考资料。

可能的原因包括：
1. 数据库中可能没有与此主题相关的文档
2. 查询措辞与文档内容不匹配
3. 相关文档的向量表示与查询向量表示的相似度不够高

您的问题是：{request.prompt}

请注意，以下回答基于通用知识而非系统中的文档内容。
"""
                else:
                    completion_prompt = f"""
You are a knowledgeable assistant, especially skilled at answering questions about specific data in documents.

I couldn't find any reference material in the document database related to "{request.prompt}". Since this appears to be a question about specific data (possibly involving rankings, tables or numeric information), I can only provide accurate answers based on actual document content.

Please consider uploading relevant documents containing the information about {request.prompt}, especially if it involves specific data points, rankings, or table information that requires precise factual knowledge.

Without access to the relevant documents, I cannot provide a specific answer to your question.
"""
        
        print(f"Generated completion prompt, length: {len(completion_prompt)}")
        completion = await gemini_service.generate_completion(completion_prompt)
        
        # If debug mode enabled, return more information
        if debug:
            return {
                "completion": completion,
                "debug_info": {
                    "original_query": request.prompt,
                    "is_chinese_query": is_chinese_query,
                    "is_table_query": is_table_query,
                    "has_ranking_terms": has_ranking_terms,
                    "force_use_documents": force_use_documents,
                    "search_query": search_query,
                    "docs_found": len(similar_docs),
                    "high_similarity_docs": len(docs_with_high_similarity) if 'docs_with_high_similarity' in locals() else 0,
                    "context_length": len(context) if context else 0,
                    "document_snippets": [
                        {
                            "content": d.get("content", "")[:150] + "...",
                            "similarity": d.get("similarity", 0),
                            "source": d.get("metadata", {}).get("source", "") or 
                                     d.get("source", "") or 
                                     d.get("metadata", {}).get("pdf_filename", "") or 
                                     "Document ID: " + str(d.get("id", "Unknown ID"))
                        } 
                        for d in similar_docs[:5]
                    ] if similar_docs else []
                }
            }
        
        # When debug is false, explicitly return response without debug_info
        # Use CompletionResponse model's field pattern to ensure proper response
        return CompletionResponse(
            completion=completion,
            debug_info=None
        ).dict(exclude_none=True)  # This excludes None values from the response
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Integration query failed: {str(e)}")

@router.get("/documents", 
           response_model=Dict[str, Any], 
           summary="Get Document List", 
           description="Get document list, supports paging and source filtering",
           response_description="Returns document list")
async def get_documents(
    db: Session = Depends(get_db),
    limit: int = Query(20, ge=1, le=100, description="Maximum number of documents to return"),
    offset: int = Query(0, ge=0, description="Page offset for paging query"),
    source: Optional[str] = Query(None, description="Filter by document source, for example specific PDF file name")
):
    """
    Get document list, supports paging and filtering
    
    Get document list from the database, supports paging and filtering by source
    
    Parameters:
        db: Database session
        limit: Maximum number of documents to return, range 1-100
        offset: Page offset for paging query
        source: Filter by document source, for example specific PDF file name
        
    Returns:
        Dict: Dictionary containing document list and total count
        
    Exceptions:
        HTTPException: If failed to get document list
    """
    try:
        # Build query conditions
        query = "SELECT id, title, doc_metadata, created_at FROM documents"
        count_query = "SELECT COUNT(*) FROM documents"
        params = {}
        
        if source:
            query += " WHERE doc_metadata::text LIKE :source"
            count_query += " WHERE doc_metadata::text LIKE :source"
            params["source"] = f"%{source}%"
            
        # Add sorting and paging
        query += " ORDER BY id DESC LIMIT :limit OFFSET :offset"
        params["limit"] = limit
        params["offset"] = offset
        
        # Execute query
        result = db.execute(text(query), params).fetchall()
        
        # Get total count
        total_count = db.execute(text(count_query), 
                               {k:v for k,v in params.items() if k != 'limit' and k != 'offset'}
                              ).scalar() or 0
        
        # Process results
        documents = []
        for row in result:
            # Safe handling doc_metadata
            try:
                metadata = row.doc_metadata
                # If it's a string, try to parse as JSON
                if isinstance(metadata, str):
                    try:
                        metadata = json.loads(metadata)
                    except json.JSONDecodeError as e:
                        print(f"Failed to parse document id={row.id} metadata: {e}")
                        metadata = {}
                # If not dictionary type, use empty dictionary
                elif not isinstance(metadata, dict):
                    print(f"Document id={row.id} metadata is not valid JSON or dictionary: {type(metadata)}")
                    metadata = {}
                
                # Delete embedding vector to reduce response size
                if "_embedding" in metadata:
                    del metadata["_embedding"]
            except Exception as e:
                print(f"Failed to handle document id={row.id} metadata: {e}")
                metadata = {}
                
            documents.append({
                "id": row.id,
                "title": row.title,
                "content": row.title,  # Use title as content return
                "metadata": metadata,
                "created_at": row.created_at.isoformat() if row.created_at else None
            })
            
        # 返回符合API规范的响应，包含documents列表和total总数
        return {
            "documents": documents,
            "total": total_count
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get document list: {str(e)}")

@router.post("/upload_pdf", 
            response_model=Dict[str, Any], 
            summary="上传和处理PDF文件", 
            description="上传PDF文件，解析内容并存储到向量数据库，支持智能分块",
            response_description="返回处理结果，包括文件名、处理的区块数量和文档ID列表")
async def upload_pdf(
    file: UploadFile = File(...),
    preserve_tables: bool = True,
    use_ocr: bool = False,
    chunking_strategy: str = Query("auto", description="分块策略: 'fixed_size', 'intelligent' 或 'auto'"),
    chunk_size: int = Query(1000, description="固定分块尺寸"),
    chunk_overlap: int = Query(200, description="固定分块重叠大小"),
    save_to_database: bool = Query(True, description="是否保存到数据库"),
    db: Session = Depends(get_db),
    vector_service: VectorService = Depends(get_vector_service),
    x_chunking_strategy: Optional[str] = Header(None, description="分块策略的header参数")
):
    """上传和处理PDF文件

    Args:
        file: 要上传的PDF文件
        preserve_tables: 保留表格结构，默认为True
        use_ocr: 是否使用OCR处理扫描文档，默认为False
        chunking_strategy: 分块策略，可选 'fixed_size', 'intelligent' 或 'auto'
        chunk_size: 固定分块尺寸，默认1000
        chunk_overlap: 固定分块重叠大小，默认200
        save_to_database: 是否保存到数据库，默认True
        x_chunking_strategy: Header中指定的分块策略，优先级高于查询参数

    Returns:
        包含提取文本的JSON响应
    """
    # 检查文件类型
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="只接受PDF文件")
    
    # 调用通用文件上传处理函数
    return await upload_file(
        file=file,
        preserve_tables=preserve_tables,
        use_ocr=use_ocr,
        chunking_strategy=chunking_strategy,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        save_to_database=save_to_database,
        db=db,
        vector_service=vector_service,
        x_chunking_strategy=x_chunking_strategy
    )

@router.post("/upload_file", 
            response_model=Dict[str, Any], 
            summary="上传和处理各种文件", 
            description="上传各种文件(PDF、TXT、DOCX、CSV等)，解析内容并存储到向量数据库，支持智能分块",
            response_description="返回处理结果，包括文件名、处理的区块数量和文档ID列表")
async def upload_file(
    file: UploadFile = File(...),
    preserve_tables: bool = True,
    use_ocr: bool = False,
    chunking_strategy: str = Query("auto", description="分块策略: 'fixed_size', 'intelligent' 或 'auto'"),
    chunk_size: int = Query(1000, description="固定分块尺寸"),
    chunk_overlap: int = Query(200, description="固定分块重叠大小"),
    save_to_database: bool = Query(True, description="是否保存到数据库"),
    db: Session = Depends(get_db),
    vector_service: VectorService = Depends(get_vector_service),
    x_chunking_strategy: Optional[str] = Header(None, description="分块策略的header参数")
):
    """上传和处理各种文件

    Args:
        file: 要上传的文件（支持PDF、TXT、DOCX、CSV等）
        preserve_tables: 保留表格结构，默认为True
        use_ocr: 是否使用OCR处理扫描文档，默认为False
        chunking_strategy: 分块策略，可选 'fixed_size', 'intelligent' 或 'auto'
        chunk_size: 固定分块尺寸，默认1000
        chunk_overlap: 固定分块重叠大小，默认200
        save_to_database: 是否保存到数据库，默认True
        x_chunking_strategy: Header中指定的分块策略，优先级高于查询参数

    Returns:
        包含提取文本的JSON响应
    """
    # 如果header中提供了chunking_strategy，则覆盖查询参数
    if x_chunking_strategy:
        chunking_strategy = x_chunking_strategy
        print(f"使用header中指定的分块策略: {chunking_strategy}")
    
    # 验证分块策略参数
    valid_strategies = ["fixed_size", "intelligent", "auto"]
    if chunking_strategy not in valid_strategies:
        raise HTTPException(status_code=400, detail=f"无效的分块策略: {chunking_strategy}，有效选项为 {', '.join(valid_strategies)}")
    
    print(f"使用分块策略: {chunking_strategy}, 保存到数据库: {save_to_database}")
    
    # 检查文件类型
    file_ext = os.path.splitext(file.filename)[1].lower()
    print(f"上传文件: {file.filename}, 扩展名: {file_ext}")
    
    # 获取文件MIME类型
    mime_type = file.content_type or ""
    print(f"文件MIME类型: {mime_type}")
    
    # 保存上传的文件到临时目录
    temp_file_path = f"/tmp/{uuid.uuid4()}{file_ext}"
    content = await file.read()
    
    with open(temp_file_path, "wb") as f:
        f.write(content)
    
    extracted_text = ""
    try:
        # 根据文件类型处理文件
        if file_ext == '.pdf':
            # PDF文件处理
            if len(content) < 4 or content[:4] != b'%PDF':
                raise HTTPException(status_code=400, detail="无效的PDF文件格式")
                
            # 使用异常处理捕获PDF读取错误
            try:
                # 首先尝试使用PyPDF2提取文本（基本提取）
                pdf_reader = PyPDF2.PdfReader(temp_file_path)
                page_count = len(pdf_reader.pages)
                print(f"成功打开PDF，页数: {page_count}")
                basic_text = ""
                for page in pdf_reader.pages:
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            basic_text += page_text + "\n\n"
                    except Exception as e:
                        print(f"处理PDF页面时出错: {e}")
                        continue
            except Exception as e:
                error_msg = str(e)
                if "EOF" in error_msg or "marker" in error_msg:
                    raise HTTPException(status_code=400, detail=f"PDF文件损坏或不完整: {error_msg}")
                else:
                    raise HTTPException(status_code=400, detail=f"无法读取PDF文件: {error_msg}")
            
            # 如果需要保留表格结构，使用pdftotext工具
            if preserve_tables:
                try:
                    # 检查pdftotext是否可用
                    try:
                        subprocess.run(["which", "pdftotext"], check=True, capture_output=True)
                    except:
                        print("pdftotext工具不可用，仅使用基本文本提取")
                        extracted_text = basic_text
                        raise Exception("pdftotext工具不可用")
                        
                    # 使用pdftotext保留布局
                    layout_text = subprocess.check_output(
                        ["pdftotext", "-layout", temp_file_path, "-"],
                        stderr=subprocess.DEVNULL
                    ).decode("utf-8", errors="ignore")
                    
                    # 如果布局提取成功且文本量合理，使用它
                    if layout_text and len(layout_text) >= len(basic_text) * 0.7:
                        extracted_text = layout_text
                        
                        # 格式化表格文本，保留列对齐
                        lines = extracted_text.split('\n')
                        formatted_lines = []
                        in_table = False
                        
                        for line in lines:
                            # 检测表格行（含有多个空格序列的行）
                            if '  ' in line and any(c.isdigit() or c.isalpha() for c in line):
                                if not in_table:
                                    formatted_lines.append("\n<TABLE>")
                                    in_table = True
                                # 替换连续空白为单个制表符以保持列分隔
                                formatted_line = re.sub(r'\s{2,}', '\t', line)
                                formatted_lines.append(formatted_line)
                            else:
                                if in_table:
                                    formatted_lines.append("</TABLE>\n")
                                    in_table = False
                                formatted_lines.append(line)
                        
                        if in_table:
                            formatted_lines.append("</TABLE>")
                            
                        extracted_text = '\n'.join(formatted_lines)
                    else:
                        extracted_text = basic_text
                except Exception as e:
                    print(f"布局保留提取失败: {e}")
                    extracted_text = basic_text
            else:
                extracted_text = basic_text
                
            # 如果使用OCR但文本提取为空，尝试OCR（需要实现）
            if use_ocr and not extracted_text.strip():
                # 这里可以添加OCR处理代码
                pass
                
        elif file_ext == '.txt':
            # 文本文件处理
            with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                extracted_text = f.read()
                
        elif file_ext == '.csv':
            # CSV文件处理
            try:
                import pandas as pd
                df = pd.read_csv(temp_file_path)
                # 转换为文本格式，保留表格结构
                extracted_text = df.to_string(index=False)
            except Exception as e:
                print(f"处理CSV文件时出错: {e}")
                # 尝试使用基本方法读取
                with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    extracted_text = f.read()
                
        elif file_ext in ['.docx', '.doc']:
            # Word文档处理
            try:
                from docx import Document
                doc = Document(temp_file_path)
                paragraphs = [p.text for p in doc.paragraphs]
                extracted_text = '\n\n'.join(paragraphs)
                
                # 如果需要保留表格
                if preserve_tables and doc.tables:
                    extracted_text += "\n\n"
                    for table in doc.tables:
                        extracted_text += "<TABLE>\n"
                        for row in table.rows:
                            row_text = '\t'.join(cell.text for cell in row.cells)
                            extracted_text += row_text + '\n'
                        extracted_text += "</TABLE>\n\n"
            except Exception as e:
                print(f"处理Word文档时出错: {e}")
                raise HTTPException(status_code=400, detail=f"无法处理Word文档: {str(e)}")
                
        elif file_ext == '.xlsx' or file_ext == '.xls':
            # Excel文件处理
            try:
                import pandas as pd
                # 读取所有工作表
                excel_file = pd.ExcelFile(temp_file_path)
                sheet_texts = []
                
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    sheet_text = f"工作表: {sheet_name}\n"
                    sheet_text += "<TABLE>\n"
                    sheet_text += df.to_string(index=False)
                    sheet_text += "\n</TABLE>\n\n"
                    sheet_texts.append(sheet_text)
                
                extracted_text = '\n\n'.join(sheet_texts)
            except Exception as e:
                print(f"处理Excel文件时出错: {e}")
                raise HTTPException(status_code=400, detail=f"无法处理Excel文件: {str(e)}")
                
        else:
            # 尝试作为纯文本文件处理
            try:
                with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    extracted_text = f.read()
                print(f"未知文件类型 {file_ext}，尝试作为纯文本处理")
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"不支持的文件类型 {file_ext}: {str(e)}")
    
    finally:
        # 清理临时文件
        try:
            os.remove(temp_file_path)
        except:
            pass
    
    if not extracted_text.strip():
        raise HTTPException(status_code=400, detail="无法从文件提取文本")
        
    # 分块处理文本
    MAX_CHUNK_SIZE = 30000  # 约30k字符
    chunks = []
    document_ids = []
    
    # 根据策略选择分块方法
    if chunking_strategy == "fixed_size":
        # 固定大小分块
        if len(extracted_text) > chunk_size:
            paragraphs = re.split(r'\n\s*\n', extracted_text)
            current_chunk = ""
            
            for para in paragraphs:
                if len(current_chunk) + len(para) < chunk_size:
                    current_chunk += para + "\n\n"
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    # 可以选择添加重叠部分
                    if chunk_overlap > 0 and current_chunk:
                        # 获取最后几个段落作为重叠部分
                        overlap_paras = current_chunk.split('\n\n')[-3:]  # 取最后3个段落
                        current_chunk = '\n\n'.join(overlap_paras) + '\n\n' + para + '\n\n'
                    else:
                        current_chunk = para + "\n\n"
            
            if current_chunk:
                chunks.append(current_chunk.strip())
        else:
            chunks = [extracted_text.strip()]
            
    elif chunking_strategy == "intelligent":
        try:
            print(f"使用智能分块策略处理文件（文件类型：{file.filename}，文本长度：{len(extracted_text)}字符）")
            file_type = file_ext.lstrip('.')
            chunks = await gemini_service.intelligent_chunking(extracted_text, file_type)
            print(f"智能分块完成，生成了{len(chunks)}个块")
            
            # 记录分块策略信息
            for i, chunk in enumerate(chunks):
                print(f"块 {i+1}/{len(chunks)}: 策略={chunk['metadata'].get('strategy', 'unknown')}, 大小={len(chunk['content'])}字符")
            
            # 不保存到数据库时，直接返回分块结果
            if not save_to_database:
                return {
                    "filename": file.filename,
                    "text_chunks": chunks,
                    "chunk_strategy": "intelligent",
                    "document_ids": []
                }
        except Exception as e:
            print(f"智能分块失败，尝试固定尺寸分块: {e}")
            # 回退到固定尺寸分块
            if len(extracted_text) > chunk_size:
                current_pos = 0
                while current_pos < len(extracted_text):
                    chunk_end = min(current_pos + chunk_size, len(extracted_text))
                    chunks.append(extracted_text[current_pos:chunk_end].strip())
                    current_pos += chunk_size - chunk_overlap
            else:
                chunks = [extracted_text.strip()]
    else:  # auto
        # 根据文件类型和内容自动选择分块策略
        if len(extracted_text) < 3000:
            # 短文本直接作为一个块
            chunks = [extracted_text.strip()]
        else:
            try:
                # 尝试智能分块
                file_type = file_ext.lstrip('.')
                chunks = await gemini_service.intelligent_chunking(extracted_text, file_type)
            except Exception as e:
                print(f"自动分块中的智能分块失败: {e}，回退到固定尺寸分块")
                # 回退到固定尺寸分块
                paragraphs = re.split(r'\n\s*\n', extracted_text)
                current_chunk = ""
                
                for para in paragraphs:
                    if len(current_chunk) + len(para) < chunk_size:
                        current_chunk += para + "\n\n"
                    else:
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                        # 可以选择添加重叠部分
                        if chunk_overlap > 0 and current_chunk:
                            # 获取最后几个段落作为重叠部分
                            overlap_paras = current_chunk.split('\n\n')[-3:]  # 取最后3个段落
                            current_chunk = '\n\n'.join(overlap_paras) + '\n\n' + para + '\n\n'
                        else:
                            current_chunk = para + "\n\n"
                
                if current_chunk:
                    chunks.append(current_chunk.strip())
    
    # 如果使用智能分块，处理已经生成的结构化块
    if chunking_strategy == "intelligent" and isinstance(chunks, list) and chunks and isinstance(chunks[0], dict):
        pass  # 已经是正确格式的块
    else:
        # 转换纯文本块为所需的字典格式
        chunks = [
            {
                "content": chunk,
                "metadata": {
                    "source": file.filename,
                    "chunk": i+1,
                    "total_chunks": len(chunks),
                    "strategy": chunking_strategy,
                    "chunk_size": chunk_size if chunking_strategy == "fixed_size" else None,
                    "overlap": chunk_overlap if chunking_strategy == "fixed_size" else None,
                    "file_type": file_ext.lstrip('.'),
                    "import_timestamp": datetime.now().strftime("%Y%m%d%H%M%S")
                }
            }
            for i, chunk in enumerate(chunks)
        ]
    
    # 保存到数据库
    if save_to_database:
        try:
            processed_chunks = []
            for chunk in chunks:
                # 处理metadata格式
                if isinstance(chunk, dict) and "content" in chunk:
                    content = chunk["content"]
                    metadata = chunk.get("metadata", {})
                else:
                    content = chunk
                    metadata = {
                        "source": file.filename,
                        "strategy": chunking_strategy
                    }
                
                # 确保metadata中包含文件名和时间戳
                if "source" not in metadata:
                    metadata["source"] = file.filename
                if "import_timestamp" not in metadata:
                    metadata["import_timestamp"] = datetime.now().strftime("%Y%m%d%H%M%S")
                
                processed_chunks.append({
                    "content": content,
                    "metadata": metadata,
                    "chunking_strategy": chunking_strategy
                })
            
            # 批量保存文档
            result = await vector_service.add_documents(db, processed_chunks)
            
            # 收集文档ID
            for doc in result:
                if hasattr(doc, 'id'):
                    document_ids.append(doc.id)
                elif isinstance(doc, dict) and "id" in doc:
                    document_ids.append(doc["id"])
        except Exception as e:
            print(f"保存文档到数据库失败: {e}")
            traceback.print_exc()
            # 返回处理结果，但提示保存失败
            return {
                "filename": file.filename,
                "text_chunks": chunks,
                "chunk_strategy": chunking_strategy,
                "document_ids": [],
                "error": f"保存到数据库失败: {str(e)}"
            }
    
    return {
        "filename": file.filename,
        "text_chunks": chunks,
        "chunk_strategy": chunking_strategy,
        "document_ids": document_ids
    }

@router.post("/analyze-documents", 
            response_model=CompletionResponse, 
            summary="Analyze Document Content", 
            description="Deeply analyze document content, extract key concepts and themes",
            response_description="Returns detailed structured document analysis result")
async def analyze_documents(
    request: QueryRequest,
    db: Session = Depends(get_db),
    vector_service: VectorService = Depends(get_vector_service),
):
    """
    Deeply analyze document content, extract key concepts and themes
    
    Retrieve related documents and use Gemini model for deep analysis, including key concepts, theme classification, and important information points
    
    Parameters:
        request: Request object containing query text and limit count
        db: Database session
        vector_service: Vector service instance
        
    Returns:
        CompletionResponse: Response object containing detailed analysis result
        
    Exceptions:
        HTTPException: If failed to analyze document
    """
    try:
        # First retrieve related documents
        results = await vector_service.search_similar(db, request.query, request.limit)
        
        if not results:
            return {"completion": "No related documents found. Please try adjusting query or increasing document library."}
        
        # Prepare context
        context = await gemini_service.prepare_context(request.query, results)
        
        # Create detailed analysis prompt
        analysis_prompt = f"""
You are a professional document analysis expert. Please perform deep analysis on the following content and extract key themes and concepts:

{context}

For query "{request.query}", please provide the following analysis:

1. Main concepts: List all key concepts mentioned in the document, each with a brief explanation.
2. Theme classification: Categorize content by theme.
3. Key points: Summarize the most important information points related to the query from the document.
4. Complete answer: Provide a comprehensive and concise answer based on document content for the user's query.

Note: If document content is truncated or incomplete, please indicate in the answer.
"""
        
        # Generate analysis
        analysis = await gemini_service.generate_completion(analysis_prompt)
        
        return {"completion": analysis}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to analyze document: {str(e)}")

@router.post("/clear-alloydb", 
            response_model=Dict[str, Any], 
            summary="Clear AlloyDB", 
            description="Clear all data tables in AlloyDB, this is a dangerous operation, need to confirm parameters",
            response_description="Returns operation result information")
async def clear_alloydb(confirmation: str = Query(..., description="Confirm string, must be 'confirm_clear_alloydb'")):
    """
    Clear all data tables in AlloyDB. Extremely cautious use!
    
    Clear all table data in AlloyDB database. To prevent accidental operations, need to provide confirmation parameter
    
    Parameters:
        confirmation: Confirm string, must be 'confirm_clear_alloydb'
        
    Returns:
        Dict: Dictionary containing operation result information
        
    Exceptions:
        HTTPException: If confirmation parameter is incorrect or clear operation fails
    """
    if confirmation != "confirm_clear_alloydb":
        raise HTTPException(status_code=400, detail="Need to confirm parameters to execute this operation")
    
    try:
        deleted_tables = 0
        table_counts = {}
        
        with engine.connect() as connection:
            # Get all table names
            inspector = inspect(engine)
            tables = inspector.get_table_names()
            
            # First delete tables not referenced (no foreign key dependencies)
            for table in tables:
                try:
                    # Get record count in table
                    count = connection.execute(text(f"SELECT COUNT(*) FROM {table}")).scalar()
                    table_counts[table] = count
                    
                    # Use DELETE instead of TRUNCATE
                    connection.execute(text(f"DELETE FROM {table}"))
                    deleted_tables += 1
                except Exception as e:
                    print(f"Failed to clear table {table}: {e}")
            
            connection.commit()
            
            return {
                "status": "success", 
                "deleted_tables": deleted_tables,
                "table_counts": table_counts,
                "timestamp": datetime.now().isoformat()
            }
    except SQLAlchemyError as e:
        raise HTTPException(status_code=500, detail=f"Failed to clear AlloyDB: {str(e)}")

@router.get("/documents/{document_id}", 
           response_model=Dict[str, Any], 
           summary="Get Single Document", 
           description="Get detailed information about a single document based on ID",
           response_description="Returns document details")
async def get_document(document_id: int = PathParam(..., description="Document ID"), db: Session = Depends(get_db)):
    """
    Get document by specified ID
    
    Get detailed information about a single document based on document ID
    
    Parameters:
        document_id: Document's unique identifier
        db: Database session
        
    Returns:
        Dict: Dictionary containing document details
        
    Exceptions:
        HTTPException: If document does not exist or retrieval fails
    """
    try:
        # Try to get document, first try id as integer
        result = db.execute(
            text("SELECT id, title, doc_metadata, created_at FROM documents WHERE id::text = :id"),
            {"id": str(document_id)}
        ).fetchone()
        
        if not result:
            # If not found, may need to try other id format
            raise HTTPException(status_code=404, detail=f"No document found with ID {document_id}")
            
        # Safe handling doc_metadata
        try:
            metadata = result.doc_metadata
            # If it's a string, try to parse as JSON
            if isinstance(metadata, str):
                try:
                    metadata = json.loads(metadata)
                except json.JSONDecodeError as e:
                    print(f"Failed to parse document id={result.id} metadata: {e}")
                    metadata = {}
            # If not dictionary type, use empty dictionary
            elif not isinstance(metadata, dict):
                print(f"Document id={result.id} metadata is not valid JSON or dictionary: {type(metadata)}")
                metadata = {}
            
            # Delete embedding vector to reduce response size
            if "_embedding" in metadata:
                del metadata["_embedding"]
        except Exception as e:
            print(f"Failed to handle document id={result.id} metadata: {e}")
            metadata = {}
            
        return {
            "id": result.id,
            "title": result.title,
            "content": result.title,  # Use title as content return
            "metadata": metadata,
            "created_at": result.created_at.isoformat() if result.created_at else None
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get document: {str(e)}")

@router.delete("/documents/{document_id}", 
             response_model=Dict[str, Any], 
             summary="Delete Document", 
             description="Delete specified document from the database",
             response_description="Returns status of the deletion operation")
async def delete_document(document_id: int = PathParam(..., description="Document ID"), db: Session = Depends(get_db)):
    """
    Delete document by ID
    
    Delete a document and its metadata from the database
    
    Parameters:
        document_id: Document's unique identifier
        db: Database session
        
    Returns:
        Dict: Status information about the deletion operation
        
    Exceptions:
        HTTPException: If document does not exist or deletion fails
    """
    try:
        # Check if document exists
        check_result = db.execute(
            text("SELECT id FROM documents WHERE id::text = :id"),
            {"id": str(document_id)}
        ).fetchone()
        
        if not check_result:
            raise HTTPException(status_code=404, detail=f"No document found with ID {document_id}")
        
        # Delete document
        delete_result = db.execute(
            text("DELETE FROM documents WHERE id::text = :id"),
            {"id": str(document_id)}
        )
        
        # Commit the transaction
        db.commit()
        
        # Return success status
        return {
            "status": "success",
            "message": f"Document {document_id} deleted successfully",
            "deleted_records": delete_result.rowcount
        }
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()  # 确保在发生错误时回滚事务
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {str(e)}")

@router.post("/upload-dual-chunking", 
            response_model=Dict[str, Any], 
            summary="上传使用双重分块策略的文档", 
            description="上传文档并同时使用固定尺寸分块和智能分块进行处理，方便后续比较不同策略的检索效果",
            response_description="返回处理结果，包括文件名、处理的区块数量和文档ID列表")
async def upload_dual_chunking(
    file: UploadFile = File(...),
    fixed_chunk_size: int = Query(1000, description="固定尺寸分块的块大小"),
    fixed_overlap: int = Query(200, description="固定尺寸分块的重叠大小"),
    db: Session = Depends(get_db),
    vector_service: VectorService = Depends(get_vector_service),
):
    """
    上传文档并同时使用固定尺寸分块和智能分块进行处理
    
    参数:
        file: 上传的文件
        fixed_chunk_size: 固定尺寸分块的块大小
        fixed_overlap: 固定尺寸分块的重叠大小
        db: 数据库会话
    
    返回:
        包含处理结果的对象
    
    异常:
        HTTPException: 如果文件处理失败
    """
    try:
        # 读取文件内容
        file_content = await file.read()
        file_extension = os.path.splitext(file.filename)[1].lower()
        
        # 文件名时间戳
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        
        # 处理PDF文件
        if file_extension == '.pdf':
            # 创建临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(file_content)
                temp_path = temp_file.name
                
            try:
                # 打开PDF
                with open(temp_path, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    num_pages = len(pdf.pages)
                    print(f"PDF文件包含 {num_pages} 页")
                    
                    # 提取PDF文本
                    full_text = ""
                    for i in range(num_pages):
                        page = pdf.pages[i]
                        page_text = page.extract_text()
                        if page_text:
                            full_text += page_text + "\n\n"
                    
                    print(f"提取的文本长度: {len(full_text)} 字符")
                    
                    # 使用固定尺寸分块
                    fixed_chunks = []
                    for i in range(0, len(full_text), fixed_chunk_size - fixed_overlap):
                        chunk_text = full_text[i:i + fixed_chunk_size]
                        if chunk_text.strip():
                            fixed_chunks.append({
                                "content": chunk_text,
                                "metadata": {
                                    "source": file.filename,
                                    "pdf_filename": file.filename,
                                    "chunk": len(fixed_chunks) + 1,
                                    "import_timestamp": timestamp,
                                    "page_range": f"PDF extraction does not track exact page mapping",
                                    "strategy": "fixed_size",
                                    "chunk_size": fixed_chunk_size,
                                    "overlap": fixed_overlap
                                }
                            })
                    
                    # 使用Gemini智能分块
                    intelligent_chunks = await gemini_service.intelligent_chunking(full_text, "pdf")
                    for i, chunk in enumerate(intelligent_chunks):
                        # 添加额外的元数据
                        chunk["metadata"].update({
                            "source": file.filename,
                            "pdf_filename": file.filename,
                            "import_timestamp": timestamp,
                            "page_range": f"PDF extraction does not track exact page mapping"
                        })
                    
                    # 存储固定尺寸分块到数据库
                    fixed_doc_ids = []
                    for chunk in fixed_chunks:
                        doc = await vector_service.add_document(
                            db, 
                            chunk["content"], 
                            chunk["metadata"],
                            "fixed_size"  # 使用固定尺寸分块策略
                        )
                        fixed_doc_ids.append(doc.id)
                    
                    # 存储智能分块到数据库
                    intelligent_doc_ids = []
                    for chunk in intelligent_chunks:
                        doc = await vector_service.add_document(
                            db, 
                            chunk["content"], 
                            chunk["metadata"],
                            "intelligent"  # 使用智能分块策略
                        )
                        intelligent_doc_ids.append(doc.id)
                    
                    # 返回处理结果
                    return {
                        "filename": file.filename,
                        "file_size": len(file_content),
                        "text_length": len(full_text),
                        "fixed_size_chunks": {
                            "count": len(fixed_chunks),
                            "chunk_size": fixed_chunk_size,
                            "overlap": fixed_overlap,
                            "doc_ids": fixed_doc_ids
                        },
                        "intelligent_chunks": {
                            "count": len(intelligent_chunks),
                            "doc_ids": intelligent_doc_ids
                        },
                        "status": "success"
                    }
            finally:
                # 删除临时文件
                os.unlink(temp_path)
        else:
            # 不支持的文件类型
            raise HTTPException(status_code=400, detail=f"不支持的文件类型：{file_extension}")
    except Exception as e:
        # 处理异常
        error_message = f"处理文件失败: {str(e)}"
        print(error_message)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=error_message)

@router.post("/compare-strategies", 
            response_model=Dict[str, Any], 
            summary="比较不同分块策略的检索效果", 
            description="对同一个查询使用不同的分块策略进行检索，并对比结果",
            response_description="返回不同策略的检索结果和对比分析")
async def compare_strategies(
    query_request: QueryRequest,
    vector_service: VectorService = Depends(get_vector_service)
):
    """对比不同的搜索策略并返回结果"""
    result = await vector_service.compare_search_strategies(
        query_request.query,
        query_request.limit
    )
    
    return result 

@router.get("/benchmark", 
            response_class=HTMLResponse,
            summary="向量搜索基准测试页面", 
            description="提供向量搜索基准测试的Web界面，可以运行测试并查看结果",
            response_description="返回基准测试页面或结果")
async def benchmark_page(
    request: Request,
    run_test: bool = Query(False, description="是否执行基准测试"),
    questions: Optional[str] = Query(None, description="测试问题，用分号分隔"),
    limit: int = Query(5, description="每个查询返回的最大结果数"),
    source_filter: Optional[str] = Query(None, description="源文档过滤条件")
):
    """
    向量搜索基准测试页面
    
    提供向量搜索基准测试的Web界面，可以运行测试并查看结果
    
    Parameters:
        request: FastAPI请求对象
        run_test: 是否执行基准测试
        questions: 测试问题，用分号分隔
        limit: 每个查询返回的最大结果数
        source_filter: 源文档过滤条件
        
    Returns:
        HTMLResponse: 基准测试页面或结果
    """
    
    # 如果不是请求运行测试，则返回测试表单页面
    if not run_test:
        return f"""
        <!DOCTYPE html>
        <html lang="zh-CN">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>向量搜索基准测试</title>
            <style>
                body {{
                    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, 'Open Sans', 'Helvetica Neue', sans-serif;
                    line-height: 1.6;
                    color: #333;
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 20px;
                    background-color: #f5f8fa;
                }}
                h1, h2, h3 {{
                    color: #0366d6;
                }}
                .container {{
                    background-color: #fff;
                    border-radius: 8px;
                    box-shadow: 0 1px 3px rgba(0,0,0,0.12);
                    padding: 20px;
                    margin-bottom: 20px;
                }}
                input, textarea, select {{
                    width: 100%;
                    padding: 10px;
                    margin-bottom: 15px;
                    border: 1px solid #ddd;
                    border-radius: 4px;
                    box-sizing: border-box;
                }}
                textarea {{
                    min-height: 150px;
                    font-family: inherit;
                }}
                label {{
                    display: block;
                    margin-bottom: 5px;
                    font-weight: 600;
                }}
                button {{
                    background-color: #0366d6;
                    color: white;
                    border: none;
                    padding: 10px 15px;
                    border-radius: 4px;
                    cursor: pointer;
                    font-weight: 500;
                }}
                button:hover {{
                    background-color: #0056b3;
                }}
                .footer {{
                    text-align: center;
                    margin-top: 30px;
                    padding: 20px;
                    color: #586069;
                    font-size: 0.9em;
                }}
                .loading {{
                    display: none;
                    text-align: center;
                    margin: 20px 0;
                }}
                .loading.active {{
                    display: block;
                }}
                .spinner {{
                    border: 4px solid #f3f3f3;
                    border-top: 4px solid #0366d6;
                    border-radius: 50%;
                    width: 30px;
                    height: 30px;
                    animation: spin 2s linear infinite;
                    margin: 0 auto;
                }}
                @keyframes spin {{
                    0% {{ transform: rotate(0deg); }}
                    100% {{ transform: rotate(360deg); }}
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <h1>向量搜索基准测试</h1>
                <p>此工具将对不同的分块策略进行基准测试，比较它们的检索性能。</p>
                
                <form id="benchmarkForm" action="/api/v1/benchmark" method="get">
                    <input type="hidden" name="run_test" value="true">
                    
                    <div>
                        <label for="questions">测试问题（每行一个问题）：</label>
                        <textarea id="questions" name="questions" placeholder="请输入测试问题，每行一个。如果留空，将使用默认问题集。"></textarea>
                    </div>
                    
                    <div>
                        <label for="limit">每个查询返回的最大结果数：</label>
                        <input type="number" id="limit" name="limit" value="5" min="1" max="20">
                    </div>
                    
                    <div>
                        <label for="source_filter">源文档过滤条件（可选）：</label>
                        <input type="text" id="source_filter" name="source_filter" placeholder="例如：document1.pdf">
                    </div>
                    
                    <button type="submit" onclick="showLoading()">运行基准测试</button>
                </form>
                
                <div id="loading" class="loading">
                    <div class="spinner"></div>
                    <p>正在运行基准测试，请稍候...</p>
                    <p>这可能需要几分钟时间，取决于测试问题的数量和系统负载。</p>
                </div>
            </div>
            
            <div class="footer">
                <p>Gemini Vector Search © {datetime.now().year}</p>
            </div>
            
            <script>
                function showLoading() {{
                    document.getElementById('benchmarkForm').style.display = 'none';
                    document.getElementById('loading').classList.add('active');
                }}
                
                // 将textarea中的多行内容转换为分号分隔的格式
                document.getElementById('benchmarkForm').addEventListener('submit', function(e) {{
                    const textarea = document.getElementById('questions');
                    if (textarea.value.trim()) {{
                        textarea.value = textarea.value.split('\\n')
                            .map(line => line.trim())
                            .filter(line => line.length > 0)
                            .join(';');
                    }}
                }});
            </script>
        </body>
        </html>
        """
    
    # 运行基准测试
    try:
        # 准备参数
        benchmark_script = os.path.join(os.getcwd(), "benchmark_search.py")
        report_dir = os.path.join(os.getcwd(), "static", "reports")
        os.makedirs(report_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = os.path.join(report_dir, f"benchmark_report_{timestamp}.html")
        
        # 解析问题列表
        question_list = []
        if questions:
            question_list = [q.strip() for q in questions.split(";") if q.strip()]
        
        # 准备命令行参数
        cmd_args = [sys.executable, benchmark_script, "--output", output_file]
        
        # 添加请求的URL
        base_url = str(request.base_url).rstrip('/')
        cmd_args.extend(["--url", base_url])
        
        # 添加查询限制
        cmd_args.extend(["--limit", str(limit)])
        
        # 添加源文档过滤条件（如果提供）
        if source_filter:
            cmd_args.extend(["--source-filter", source_filter])
        
        # 添加问题列表（如果提供）
        if question_list:
            cmd_args.extend(["--questions"] + question_list)
        
        # 运行基准测试
        print(f"执行基准测试命令: {' '.join(cmd_args)}")
        process = subprocess.run(
            cmd_args,
            capture_output=True,
            text=True,
            check=True
        )
        
        # 读取生成的报告
        with open(output_file, 'r', encoding='utf-8') as f:
            report_content = f.read()
        
        return HTMLResponse(content=report_content)
    
    except Exception as e:
        error_message = f"运行基准测试时出错: {str(e)}"
        if hasattr(e, 'stdout') and e.stdout:
            error_message += f"\n\n标准输出:\n{e.stdout}"
        if hasattr(e, 'stderr') and e.stderr:
            error_message += f"\n\n错误输出:\n{e.stderr}"
        
        return HTMLResponse(
            content=f"""
            <!DOCTYPE html>
            <html lang="zh-CN">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>基准测试错误</title>
                <style>
                    body {{
                        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, 'Open Sans', 'Helvetica Neue', sans-serif;
                        line-height: 1.6;
                        color: #333;
                        max-width: 800px;
                        margin: 0 auto;
                        padding: 20px;
                        background-color: #f5f8fa;
                    }}
                    h1 {{
                        color: #d73a49;
                    }}
                    .error-container {{
                        background-color: #fff;
                        border-radius: 8px;
                        box-shadow: 0 1px 3px rgba(0,0,0,0.12);
                        padding: 20px;
                        margin-bottom: 20px;
                        border-left: 5px solid #d73a49;
                    }}
                    .error-details {{
                        background-color: #f6f8fa;
                        border: 1px solid #e1e4e8;
                        border-radius: 4px;
                        padding: 15px;
                        overflow: auto;
                        white-space: pre-wrap;
                        font-family: monospace;
                    }}
                    .back-button {{
                        background-color: #0366d6;
                        color: white;
                        border: none;
                        padding: 10px 15px;
                        border-radius: 4px;
                        cursor: pointer;
                        font-weight: 500;
                        text-decoration: none;
                        display: inline-block;
                        margin-top: 20px;
                    }}
                </style>
            </head>
            <body>
                <div class="error-container">
                    <h1>基准测试运行失败</h1>
                    <p>在运行基准测试时遇到错误：</p>
                    
                    <div class="error-details">
                        {error_message}
                    </div>
                    
                    <a href="/api/v1/benchmark" class="back-button">返回</a>
                </div>
            </body>
            </html>
            """,
            status_code=500
        ) 
